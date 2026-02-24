"""Unit tests for WordParser (.docx)."""

from pathlib import Path

from docx import Document

from app.parsers.base import ParserError
from app.parsers.word_parser import WordParser


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")
    document.save(path)


def test_word_parser_extracts_paragraph_text(tmp_path: Path) -> None:
    """Word parser should return text that contains inserted paragraph lines."""

    path = tmp_path / "resume.docx"
    _create_docx(path)

    parser = WordParser()
    parsed = parser.parse(str(path))

    assert "Jane Doe" in parsed
    assert "jane@example.com" in parsed


def test_word_parser_extracts_table_text(tmp_path: Path) -> None:
    """Word parser should include table cell text in extracted output."""

    path = tmp_path / "resume.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "LLM"
    document.save(path)

    parser = WordParser()
    parsed = parser.parse(str(path))

    assert "Python" in parsed
    assert "LLM" in parsed


def test_word_parser_can_parse_docx_suffix_case_insensitive() -> None:
    """Word parser should match suffix regardless of casing."""

    parser = WordParser()

    assert parser.can_parse("resume.DOCX")


def test_word_parser_raises_parser_error_for_missing_file() -> None:
    """Word parser should raise ParserError when the target file cannot be opened."""

    parser = WordParser()

    try:
        parser.parse("does_not_exist.docx")
    except ParserError as exc:
        assert "does_not_exist.docx" in str(exc)
    else:
        raise AssertionError("Expected ParserError for missing file")
