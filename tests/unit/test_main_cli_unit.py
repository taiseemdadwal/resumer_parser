"""Unit-level CLI tests that execute ``main(argv)`` directly for coverage."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app import main as main_module


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Email: jane@example.com")
    document.add_paragraph("Skill Set: Python, AWS, Docker, Kubernetes")
    document.save(path)


def test_main_returns_0_for_valid_docx(capsys, tmp_path: Path) -> None:
    """CLI main should return success and print JSON for a valid DOCX input."""

    path = tmp_path / "resume.docx"
    _create_docx(path)

    code = main_module.main(
        ["--file", str(path), "--name-mode", "rule", "--email-mode", "regex", "--skills-mode", "fake"]
    )

    out = capsys.readouterr().out
    assert code == 0
    assert '"name": "Jane Doe"' in out
    assert '"Python"' in out


def test_main_returns_0_for_stub_mode(capsys, tmp_path: Path) -> None:
    """CLI main should support stub skills mode and emit an empty skills list."""

    path = tmp_path / "resume.docx"
    _create_docx(path)

    code = main_module.main(
        ["--file", str(path), "--name-mode", "stub", "--email-mode", "stub", "--skills-mode", "stub"]
    )

    out = capsys.readouterr().out
    assert code == 0
    assert '"name": null' in out
    assert '"email": null' in out
    assert '"skills": []' in out


def test_main_returns_2_for_argument_errors() -> None:
    """CLI main should map argparse failures to the user-error exit code."""

    code = main_module.main([])

    assert code == 2


def test_main_returns_1_for_parser_errors(capsys, tmp_path: Path) -> None:
    """CLI main should map parser failures to runtime error exit code 1."""

    path = tmp_path / "resume.txt"
    path.write_text("Jane Doe", encoding="utf-8")

    code = main_module.main(["--file", str(path)])

    err = capsys.readouterr().err
    assert code == 1
    assert "ParserError" in err


def test_main_returns_1_for_llm_mode_without_api_key(capsys, monkeypatch, tmp_path: Path) -> None:
    """CLI main should return runtime-error code when llm mode lacks API key."""

    path = tmp_path / "resume.docx"
    _create_docx(path)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    code = main_module.main(["--file", str(path), "--skills-mode", "llm"])

    err = capsys.readouterr().err
    assert code == 1
    assert "ValueError" in err
