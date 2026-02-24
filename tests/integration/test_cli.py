"""CLI behavior tests for ``python -m app.main``."""

from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

from docx import Document


def _create_docx(path: Path) -> None:
    """Create a small DOCX file used by CLI integration tests."""

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Email: jane@example.com")
    document.save(path)


def _cli_env() -> dict[str, str]:
    """Build environment with ``src`` on ``PYTHONPATH`` for subprocess execution."""

    env = os.environ.copy()
    root = Path(__file__).resolve().parents[1]
    src = str(root / "src")
    existing = env.get("PYTHONPATH")
    env["PYTHONPATH"] = src if not existing else f"{src}{os.pathsep}{existing}"
    return env


def test_cli_parses_docx_and_prints_json(tmp_path: Path) -> None:
    """CLI should parse a DOCX file and emit valid JSON to stdout."""

    path = tmp_path / "resume.docx"
    _create_docx(path)

    completed = subprocess.run(
        [sys.executable, "-m", "app.main", "--file", str(path)],
        capture_output=True,
        text=True,
        check=False,
        env=_cli_env(),
    )

    assert completed.returncode == 0
    payload = json.loads(completed.stdout)
    assert payload["name"] == "Jane Doe"
    assert payload["email"] == "jane@example.com"
    assert "skills" in payload


def test_cli_unknown_extension_returns_error(tmp_path: Path) -> None:
    """CLI should return code 1 for unsupported file suffixes."""

    path = tmp_path / "resume.txt"
    path.write_text("Jane Doe", encoding="utf-8")

    completed = subprocess.run(
        [sys.executable, "-m", "app.main", "--file", str(path)],
        capture_output=True,
        text=True,
        check=False,
        env=_cli_env(),
    )

    assert completed.returncode == 1
    assert "ParserError" in completed.stderr
