"""Integration tests for ResumeParserFramework orchestration."""

from pathlib import Path

from docx import Document

from app.extractors.email_extractor import EmailExtractor
from app.extractors.name_extractor import NameExtractor
from app.extractors.skills_extractor import LLMClient, SkillsExtractor
from app.framework import ResumeParserFramework
from app.parsers.base import ParserError
from app.parsers.pdf_parser import PDFParser
from app.parsers.word_parser import WordParser
from app.resume_extractor import ResumeExtractor


class FakeLLMClient(LLMClient):
    """Offline fake client used in integration testing."""

    def complete(self, prompt: str) -> str:
        _ = prompt
        return '["Python", "LLM"]'


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Email: jane@example.com")
    document.save(path)


def test_framework_parse_resume_end_to_end_with_word_file(tmp_path: Path) -> None:
    """Framework should parse file text then extract all required fields."""

    path = tmp_path / "resume.docx"
    _create_docx(path)

    resume_extractor = ResumeExtractor(
        {
            "name": NameExtractor(),
            "email": EmailExtractor(),
            "skills": SkillsExtractor(FakeLLMClient()),
        }
    )
    framework = ResumeParserFramework([PDFParser(), WordParser()], resume_extractor)

    result = framework.parse_resume(str(path))

    assert result.name == "Jane Doe"
    assert result.email == "jane@example.com"
    assert result.skills == ["Python", "LLM"]


def test_framework_raises_parser_error_for_unknown_suffix() -> None:
    """Framework should fail with ParserError when no parser supports the suffix."""

    extractor = ResumeExtractor(
        {
            "name": NameExtractor(),
            "email": EmailExtractor(),
            "skills": SkillsExtractor(FakeLLMClient()),
        }
    )
    framework = ResumeParserFramework([PDFParser(), WordParser()], extractor)

    try:
        framework.parse_resume("resume.txt")
    except ParserError as exc:
        assert "resume.txt" in str(exc)
    else:
        raise AssertionError("Expected ParserError for unsupported suffix")


def test_resume_extractor_requires_all_field_keys() -> None:
    """Coordinator should raise ValueError when required extractor keys are missing."""

    try:
        ResumeExtractor({"name": NameExtractor(), "email": EmailExtractor()})
    except ValueError as exc:
        assert "skills" in str(exc)
    else:
        raise AssertionError("Expected ValueError for missing 'skills' extractor")
