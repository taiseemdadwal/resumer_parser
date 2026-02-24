"""DOCX parser implementation using ``python-docx``."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import FileParser, ParserError


class WordParser(FileParser):
    """Extract text from DOCX files.

    Paragraphs and table cells are included in output so resumes that use simple table
    layouts still produce useful text for extraction.
    """

    def can_parse(self, file_path: str) -> bool:
        """Check whether this parser supports the provided file path.

        This method performs case-insensitive suffix matching for DOCX files.

        Args:
            file_path: Candidate file path.

        Returns:
            bool: ``True`` when path suffix is ``.docx`` (case-insensitive).

        Raises:
            None.
        """

        return Path(file_path).suffix.lower() == ".docx"

    def parse(self, file_path: str) -> str:
        """Parse DOCX content into newline-separated plain text.

        Paragraph text and table cell text are both included in the output.

        Args:
            file_path: Path to a Word DOCX file.

        Returns:
            str: Plain-text content assembled from paragraphs and table cells.

        Raises:
            ParserError: If the DOCX cannot be opened or parsed.
        """

        try:
            document = Document(file_path)
        except (OSError, PackageNotFoundError, ValueError) as exc:
            raise ParserError(f"Failed to parse Word file '{file_path}': {exc}") from exc

        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)

        return "\n".join(lines)
